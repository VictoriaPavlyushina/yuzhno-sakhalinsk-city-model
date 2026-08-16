from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path(r"C:\Users\victo\Downloads\Сводная_справка_Южно_Сахалинск_2016_2025_13082026.docx")
OUTPUT = Path(r"C:\Users\victo\Documents\Codex\2026-08-11\new-chat\public\Сводная_справка_Южно-Сахалинск_2016-2025.docx")

DARK = "123536"
TEAL = "73C8C6"
WARM = "E5A06F"
MUTED = "647878"
PALE = "EAF3EF"
WHITE = "FFFFFF"


def snapshot(doc):
    return (
        [p.text for p in doc.paragraphs],
        [[[p.text for p in cell.paragraphs] for cell in row.cells] for table in doc.tables for row in table.rows],
    )


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_bottom_border(paragraph, color, size=18):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


doc = Document(SOURCE)
original = snapshot(doc)

section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.6)
section.left_margin = Cm(2.1)
section.right_margin = Cm(1.8)
section.header_distance = Cm(0.7)
section.footer_distance = Cm(0.7)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(9.5)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.14

for style_name in ("Normal (Web)", "List Paragraph", "List Number"):
    if style_name in styles:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(9.5)
        style.font.color.rgb = RGBColor.from_string(DARK)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.14

heading_specs = {
    "Heading 1": (30, DARK, 18, 12),
    "Heading 2": (18, DARK, 18, 8),
    "Heading 3": (12, TEAL, 12, 5),
    "Подраздел": (11, TEAL, 10, 4),
}
for name, (size, color, before, after) in heading_specs.items():
    if name not in styles:
        continue
    style = styles[name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

if "Intense Quote" in styles:
    quote = styles["Intense Quote"]
    quote.font.name = "Arial"
    quote.font.size = Pt(9)
    quote.font.italic = False
    quote.font.color.rgb = RGBColor.from_string(DARK)
    quote.paragraph_format.left_indent = Cm(0.6)
    quote.paragraph_format.right_indent = Cm(0.2)
    quote.paragraph_format.space_before = Pt(7)
    quote.paragraph_format.space_after = Pt(8)

for index, paragraph in enumerate(doc.paragraphs):
    if not paragraph.text.strip():
        continue
    style_name = paragraph.style.name
    if index == 0:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(150)
        paragraph.paragraph_format.space_after = Pt(20)
        paragraph.paragraph_format.keep_with_next = False
        paragraph.paragraph_format.page_break_after = True
        set_bottom_border(paragraph, TEAL, 26)
    elif style_name == "Heading 2":
        paragraph.paragraph_format.page_break_before = True
        set_bottom_border(paragraph, TEAL, 10)
    elif style_name == "Heading 3":
        set_bottom_border(paragraph, "BFD8D2", 5)
    elif style_name == "Intense Quote":
        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), PALE)
        p_pr.append(shd)
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "22")
        left.set(qn("w:color"), WARM)
        borders.append(left)
        p_pr.append(borders)
    for run in paragraph.runs:
        run.font.name = "Arial"

implicit_headings = {
    "Методология градоэкологического каркаса": "Heading 2",
    "Ключевые выводы": "Heading 2",
    "Общее описание влияния травматизма в детстве на социально-экономическое развитие": "Heading 3",
    "Основные подходы к оценке травматизма": "Heading 3",
    "Результаты расчётов потерь: основные цифры и аналитический вывод": "Heading 3",
}
for paragraph in doc.paragraphs:
    target_style = implicit_headings.get(paragraph.text.strip())
    if target_style:
        paragraph.style = styles[target_style]
        if target_style == "Heading 2":
            paragraph.paragraph_format.page_break_before = True
            set_bottom_border(paragraph, TEAL, 10)
        else:
            set_bottom_border(paragraph, "BFD8D2", 5)

for table in doc.tables:
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if row_index == 0:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, DARK)
            elif row_index % 2:
                set_cell_shading(cell, PALE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7.5)
                    if row_index == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)
                    else:
                        run.font.color.rgb = RGBColor.from_string(DARK)

header = section.header
hp = header.paragraphs[0]
hp.text = "ЮЖНО-САХАЛИНСК  ·  АНАЛИТИЧЕСКАЯ СПРАВКА  ·  2016–2025"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in hp.runs:
    run.font.name = "Arial"
    run.font.size = Pt(7)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
set_bottom_border(hp, TEAL, 7)

footer = section.footer
fp = footer.paragraphs[0]
set_bottom_border(fp, "BFD8D2", 4)
add_page_field(fp)

doc.core_properties.title = doc.paragraphs[0].text
doc.core_properties.subject = "Сводная аналитическая справка"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)

check = Document(OUTPUT)
assert snapshot(check) == original, "Document text or table content changed during styling"
print(OUTPUT)
