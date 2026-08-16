from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path(r"C:\Users\victo\Downloads\Обновление_по_новой_трассировке_в_расчете_стоимости_стойки.docx")
OUTPUT = Path(r"C:\Users\victo\Documents\Codex\2026-08-11\new-chat\outputs\019ff12b-data-package\Обновление_по_новой_трассировке.docx")

DARK = "123536"
TEAL = "73C8C6"
WARM = "E5A06F"
PALE = "EAF3EF"
INK = "17393A"
WHITE = "FFFFFF"


def snapshot(doc):
    return [p.text for p in doc.paragraphs]


def shade(paragraph, color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)


def left_border(paragraph, color, size=18):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), color)
    borders.append(left)
    p_pr.append(borders)


doc = Document(SOURCE)
original = snapshot(doc)
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.0)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.line_spacing = 1.18
normal.paragraph_format.space_after = Pt(10)

for i, paragraph in enumerate(doc.paragraphs):
    if not paragraph.text.strip():
        continue
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_after = Pt(11)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(INK)
    if i == 0:
        shade(paragraph, DARK)
        left_border(paragraph, TEAL, 26)
        paragraph.paragraph_format.space_before = Pt(90)
        paragraph.paragraph_format.space_after = Pt(22)
        paragraph.paragraph_format.left_indent = Cm(0.35)
        paragraph.paragraph_format.right_indent = Cm(0.25)
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(WHITE)
            if run.bold:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor.from_string(TEAL)
    elif i == 1:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(14)
        for run in paragraph.runs:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(DARK)
        left_border(paragraph, TEAL, 18)
    elif i in (2, 3):
        shade(paragraph, PALE)
        left_border(paragraph, TEAL, 10)
        paragraph.paragraph_format.left_indent = Cm(0.25)
        paragraph.paragraph_format.right_indent = Cm(0.2)
    elif i == 4:
        shade(paragraph, "F8EDE5")
        left_border(paragraph, WARM, 18)
        paragraph.paragraph_format.left_indent = Cm(0.25)
        paragraph.paragraph_format.right_indent = Cm(0.2)
        for run in paragraph.runs:
            if run.bold:
                run.font.color.rgb = RGBColor.from_string("8A4B2C")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
assert snapshot(Document(OUTPUT)) == original, "Document content changed during styling"
print(OUTPUT)
